# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:\Users\asus\AIS视觉融合船舶碰撞风险预警-项目书.md"
DST = r"C:\Users\asus\OneDrive\桌面\AIS视觉融合船舶碰撞风险预警-项目书-v2.docx"

doc = Document()

# ---------- 全局默认字体 ----------
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(11)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.3

def _set_font(run, latin='Times New Roman', east='宋体', size=11, bold=False, italic=False, color=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:eastAsia'), east)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    return run

def add_rich(p, text):
    """解析 **粗体** 与 `行内代码`，其余作为普通文本。"""
    for tok in re.split(r'(\*\*.*?\*\*|`[^`]*`)', text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            _set_font(p.add_run(tok[2:-2]), bold=True)
        elif tok.startswith('`') and tok.endswith('`'):
            _set_font(p.add_run(tok[1:-1]), latin='Consolas', east='宋体', size=10)
        else:
            _set_font(p.add_run(tok))

def add_heading(text, level):
    p = doc.add_heading('', level=level)
    _set_font(p.add_run(text), east='微软雅黑', size={0:22,1:16,2:14,3:12}.get(level,12), bold=True, color=(0,0,0))
    return p

def add_code_block(lines):
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        _set_font(p.add_run(ln if ln else ' '), latin='Consolas', east='宋体', size=9)

def add_table(rows):
    # rows: list of list of cell strings
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            txt = row[j] if j < len(row) else ''
            add_rich(p, txt)
            # 表头加粗 + 底纹
            if i == 0:
                for run in p.runs:
                    run.bold = True
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), 'D9E2F3')
                cell._tc.get_or_add_tcPr().append(shd)
    return table

def parse_table(header_line, data_lines):
    def split_row(line):
        line = line.strip().strip('|')
        return [c.strip() for c in line.split('|')]
    rows = [split_row(header_line)]
    for ln in data_lines:
        if re.match(r'^\s*\|[\s:|-]+\|\s*$', ln):
            continue  # 分隔行
        rows.append(split_row(ln))
    return rows

with open(SRC, encoding='utf-8') as f:
    lines = f.read().split('\n')

i = 0
N = len(lines)
while i < N:
    line = lines[i]
    stripped = line.strip()

    # 空行 / 分隔线
    if stripped == '':
        i += 1
        continue
    if re.match(r'^-{3,}$', stripped):
        i += 1
        continue

    # 代码块
    if stripped.startswith('```'):
        i += 1
        block = []
        while i < N and not lines[i].strip().startswith('```'):
            block.append(lines[i])
            i += 1
        i += 1  # 跳过结尾 ```
        add_code_block(block)
        continue

    # 标题
    if stripped.startswith('### '):
        add_heading(stripped[4:], 3)
        i += 1
        continue
    if stripped.startswith('## '):
        add_heading(stripped[3:], 2)
        i += 1
        continue
    if stripped.startswith('# '):
        p = doc.add_heading('', level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(stripped[2:]), east='微软雅黑', size=22, bold=True, color=(0,0,0))
        i += 1
        continue

    # 表格
    if stripped.startswith('|'):
        j = i
        table_lines = []
        while j < N and lines[j].strip().startswith('|'):
            table_lines.append(lines[j])
            j += 1
        rows = parse_table(table_lines[0], table_lines[1:])
        add_table(rows)
        doc.add_paragraph()  # 表后空一行
        i = j
        continue

    # 引用块
    if stripped.startswith('>'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        _set_font(p.add_run(stripped.lstrip('>').strip()), italic=True, color=(0x60,0x60,0x60))
        i += 1
        continue

    # 列表（有序/无序）
    m = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)$', line)
    if m:
        indent, marker, content = m.groups()
        is_numbered = bool(re.match(r'\d+\.', marker))
        style = 'List Number' if is_numbered else 'List Bullet'
        p = doc.add_paragraph(style=style)
        if indent:
            p.paragraph_format.left_indent = Cm(0.75)
        add_rich(p, content)
        i += 1
        continue

    # 普通段落
    p = doc.add_paragraph()
    add_rich(p, stripped)
    i += 1

doc.save(DST)
print("SAVED:", DST)
